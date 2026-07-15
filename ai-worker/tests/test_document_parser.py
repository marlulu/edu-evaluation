from pathlib import Path

import pytest
from docx import Document

from app.modules.document_validation.parser import (
    VISION_IMAGE_MAX_EDGE,
    _representative_page_indexes,
    _save_vision_jpeg,
    DocumentParseError,
    extract_docx_images,
    parse_document,
)


def test_vision_jpeg_is_resized_and_compressed(tmp_path: Path) -> None:
    from PIL import Image

    source = Image.new("RGBA", (3200, 2400), (200, 40, 60, 180))
    output = tmp_path / "evidence.jpg"

    _save_vision_jpeg(source, output)

    with Image.open(output) as compressed:
        assert max(compressed.size) == VISION_IMAGE_MAX_EDGE
        assert compressed.mode == "RGB"
    assert output.stat().st_size < 500_000


def test_pdf_page_sampling_covers_first_middle_and_last_pages() -> None:
    indexes = _representative_page_indexes(30, 12)

    assert len(indexes) == 12
    assert indexes[0] == 0
    assert indexes[-1] == 29
    assert indexes == sorted(set(indexes))


def test_parse_docx_preserves_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("作品说明", level=1)
    document.add_paragraph("这是正文。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "主题"
    table.cell(1, 1).text = "测试"
    document.save(path)

    result = parse_document(path, path.name, path.stat().st_size)

    assert result.parser == "python-docx"
    assert result.headings[0].text == "作品说明"
    assert result.tables[0].rows[1] == ["主题", "测试"]
    assert "这是正文。" in result.text
    assert result.statistics.tables == 1


def test_parse_markdown_extracts_headings(tmp_path: Path) -> None:
    path = tmp_path / "sample.md"
    path.write_text("# 标题\n\n正文", encoding="utf-8")

    result = parse_document(path, path.name, path.stat().st_size)

    assert result.headings[0].level == 1
    assert result.headings[0].text == "标题"


def test_extract_docx_images_returns_embedded_raster_images(tmp_path: Path) -> None:
    from PIL import Image

    image_path = tmp_path / "embedded.png"
    Image.new("RGB", (320, 180), (30, 90, 180)).save(image_path, "PNG")
    document_path = tmp_path / "with-image.docx"
    document = Document()
    document.add_paragraph("作品说明")
    document.add_picture(str(image_path))
    document.save(document_path)

    images = extract_docx_images(document_path, tmp_path / "images")

    assert len(images) == 1
    assert images[0].name == "image_01.jpg"
    with Image.open(images[0]) as extracted:
        assert extracted.format == "JPEG"
        assert extracted.size == (320, 180)


def test_rejects_legacy_doc(tmp_path: Path) -> None:
    path = tmp_path / "sample.doc"
    path.write_bytes(b"legacy")

    with pytest.raises(DocumentParseError) as error:
        parse_document(path, path.name, path.stat().st_size)

    assert error.value.code == "UNSUPPORTED_FORMAT"
